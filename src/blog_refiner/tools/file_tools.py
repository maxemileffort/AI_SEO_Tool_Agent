from crewai_tools import tool
import os

@tool
def read_file(file_path: str) -> str:
    """
    Reads content from DOCX, PDF, or TXT files.
    """
    _, ext = os.path.splitext(file_path)
    if ext.lower() == '.txt':
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    elif ext.lower() == '.docx':
        import docx
        doc = docx.Document(file_path)
        return '\n'.join([para.text for para in doc.paragraphs])
    elif ext.lower() == '.pdf':
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        return ''.join([page.extract_text() for page in reader.pages])
    else:
        return f"Unsupported file format: {ext}"

@tool
def write_file(file_path: str, content: str) -> str:
    """
    Writes content to DOCX, PDF, or TXT files in the ../outputs folder.
    """
    import os
    _, ext = os.path.splitext(file_path)
    file_name = os.path.basename(file_path)
    
    # Ensure the outputs folder exists
    outputs_folder = os.path.abspath(os.path.join(os.path.dirname(__file__), '../../outputs'))
    os.makedirs(outputs_folder, exist_ok=True)
    
    # Create the full path for the output file
    output_dest = os.path.join(outputs_folder, file_name)
    
    if ext.lower() == '.txt':
        with open(output_dest, 'w', encoding='utf-8') as file:
            file.write(content)
    elif ext.lower() == '.docx':
        import docx
        doc = docx.Document()
        doc.add_paragraph(content)
        doc.save(output_dest)
    elif ext.lower() == '.pdf':
        from fpdf import FPDF
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font("Arial", size=12)
        pdf.multi_cell(0, 10, content)
        pdf.output(output_dest)
    else:
        return f"Unsupported file format: {ext}"
    
    return f"File written successfully to {output_dest}"

